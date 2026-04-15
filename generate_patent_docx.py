import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_patent_docx(input_file, output_file):
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found.")
        return

    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'  # Chinese font
    font.size = Pt(10.5)  # size 5 (五号)

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        if line.startswith("# "):
            # Title
            p = doc.add_paragraph(line[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(16)  # Large title
        elif line.startswith("## "):
            # Section title
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            # Subsection title
            doc.add_heading(line[4:], level=2)
        elif line.startswith("* "):
            # Bullet point
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. "):
            # Numbered list
            doc.add_paragraph(line, style='List Number')
        else:
            # Regular text
            doc.add_paragraph(line)

    doc.save(output_file)
    print(f"Patent application saved to {output_file}")

if __name__ == "__main__":
    input_md = "patent_draft.md"
    output_docx = "patent_application.docx"
    create_patent_docx(input_md, output_docx)
